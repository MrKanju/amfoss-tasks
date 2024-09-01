import requests
import csv
from io import BytesIO, StringIO
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackQueryHandler, ContextTypes, ConversationHandler
from docx import Document

API_TOKEN = "7084452444:AAFz_ut69wnsndnldzHc4AtOOzDIEl66M4U"
BOOKS_API_KEY = "AIzaSyAu7QF8blrovpt9R7T-41VjZI-LTICo2Xs"
BOOKS_DOC = "books_collection.docx"

ADDING_BOOK = 1
DELETING_BOOK = 2

def initialize_document():
    document = Document()
    document.add_heading('Books Collection', level=1)
    document.save(BOOKS_DOC)

def append_to_document(title, url):
    document = Document(BOOKS_DOC)
    document.add_paragraph(f"{title} - {url}")
    document.save(BOOKS_DOC)

def delete_from_document(title):
    document = Document(BOOKS_DOC)
    paragraphs = [p.text for p in document.paragraphs if title not in p.text]
    document._body.clear_content()
    document.add_heading('Books Collection', level=1)
    for para in paragraphs:
        document.add_paragraph(para)
    document.save(BOOKS_DOC)

async def handle_query(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    action = query.data
    if action == 'add':
        await query.edit_message_text("Please provide the book title and URL in the format: 'Title, URL'")
        context.user_data['action'] = 'add'
    elif action == 'remove':
        await query.edit_message_text("Please provide the title of the book you want to remove:")
        context.user_data['action'] = 'remove'
    elif action == 'show':
        with open(BOOKS_DOC, 'rb') as file:
            await query.message.reply_document(document=file)

async def process_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    action = context.user_data.get('action')

    if action == 'add':
        try:
            title, url = update.message.text.split(',')
            append_to_document(title.strip(), url.strip())
            await update.message.reply_text("Book successfully added to your collection!")
        except ValueError:
            await update.message.reply_text("Format error. Use 'Title, URL'")
    elif action == 'remove':
        title = update.message.text.strip()
        delete_from_document(title)
        await update.message.reply_text("Book successfully removed from your collection!")

    context.user_data['action'] = None

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Welcome to the Book Manager Bot! Use /help for commands.")

async def request_genre(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Enter the genre you want to explore.')
    return ADDING_BOOK

async def search_books(update: Update, context: ContextTypes.DEFAULT_TYPE):
    genre = update.message.text
    books = find_books(genre)
    if not books:
        await update.message.reply_text(f"No books found for genre: {genre}. Try another genre.")
        return ConversationHandler.END
    
    csv_file = create_csv(books)
    await context.bot.send_document(chat_id=update.message.chat_id, document=csv_file, filename=f"{genre}_books.csv")
    return ConversationHandler.END

def find_books(genre):
    url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={BOOKS_API_KEY}"
    response = requests.get(url)
    if response.status_code != 200:
        print(f"Error: Received status code {response.status_code}")
        return []
    
    data = response.json()
    books_list = []
    for item in data.get('items', []):
        info = item.get('volumeInfo', {})
        books_list.append({
            "title": info.get("title", "No Title"),
            "authors": ", ".join(info.get("authors", ["Unknown"])),
            "description": info.get("description", "No Description"),
            "published_date": info.get("publishedDate", "Unknown"),
            "language": info.get("language", "Unknown"),
            "preview_link": info.get("previewLink", "No Preview Link")
        })
    return books_list

def create_csv(books_list):
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=["title", "authors", "description", "published_date", "language", "preview_link"])
    writer.writeheader()
    writer.writerows(books_list)
    output.seek(0)
    return BytesIO(output.getvalue().encode())

async def ask_for_book(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Type the book title to get a preview link.')
    return DELETING_BOOK

async def provide_preview(update: Update, context: ContextTypes.DEFAULT_TYPE):
    title = update.message.text
    preview_link = get_preview_link(title)
    
    if preview_link:
        response_text = f"Preview for <b>{title}</b>: <a href='{preview_link}'>{preview_link}</a>"
        await update.message.reply_text(response_text, parse_mode='HTML')
    else:
        await update.message.reply_text(f"No preview available for: {title}. Try another title.")
    
    return ConversationHandler.END

def get_preview_link(title):
    url = f"https://www.googleapis.com/books/v1/volumes?q=intitle:{title}&key={BOOKS_API_KEY}"
    response = requests.get(url)
    if response.status_code != 200:
        print(f"Error: Received status code {response.status_code}")
        return None
    
    data = response.json()
    if 'items' in data and data['items']:
        volume_info = data['items'][0].get('volumeInfo', {})
        return volume_info.get('previewLink', "No Preview Link")
    
    return None

async def display_options(update: Update, context: ContextTypes.DEFAULT_TYPE):
    options = [
        [InlineKeyboardButton("Add Book", callback_data='add')],
        [InlineKeyboardButton("Remove Book", callback_data='remove')],
        [InlineKeyboardButton("Show Collection", callback_data='show')]
    ]
    markup = InlineKeyboardMarkup(options)
    await update.message.reply_text('Select an action to manage your book collection:', reply_markup=markup)

async def help_info(update: Update, context: ContextTypes.DEFAULT_TYPE):
    help_message = (
        "/start - Start the bot\n"
        "/book - Search for books by genre\n"
        "/preview - Get a preview link for a book\n"
        "/manage - Manage your book collection\n"
        "/help - List of available commands"
    )
    await update.message.reply_text(help_message)

def main():
    initialize_document()

    bot = Application.builder().token(API_TOKEN).build()

    book_search_handler = ConversationHandler(
        entry_points=[CommandHandler('book', request_genre)],
        states={
            ADDING_BOOK: [MessageHandler(filters.TEXT & ~filters.COMMAND, search_books)],
        },
        fallbacks=[]
    )

    preview_handler = ConversationHandler(
        entry_points=[CommandHandler('preview', ask_for_book)],
        states={
            DELETING_BOOK: [MessageHandler(filters.TEXT & ~filters.COMMAND, provide_preview)],
        },
        fallbacks=[]
    )

    bot.add_handler(CommandHandler("start", start))
    bot.add_handler(book_search_handler)
    bot.add_handler(preview_handler)
    bot.add_handler(CommandHandler("manage", display_options))
    bot.add_handler(CallbackQueryHandler(handle_query))
    bot.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, process_input))
    bot.add_handler(CommandHandler("help", help_info))

    print("Bot is active and running!")
    bot.run_polling()

if __name__ == "__main__":
    main()

